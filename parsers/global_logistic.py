"""Парсер «Global Logistic» — DOCX FTL/TIR тарифы Китай → Россия."""

from __future__ import annotations

import re
from pathlib import Path

import docx

from .models import TariffSegment
from .utils import to_segments as _to_segments
from shared import get_country, get_country_for_port

_COMPANY = "Global Logistic"

_LOADING_CITY_MAP = {
    "Tianjin/Qingdao": ("Тяньцзинь", "Китай"),
    "Shanghai/Ningbo": ("Шанхай", "Китай"),
    "Shenzhen/Guangzhou": ("Шэньчжэнь", "Китай"),
    "Tianjin/Beijing": ("Тяньцзинь", "Китай"),
}

_DESTINATION_MAP = {
    "Moscow": ("Москва", "Россия"),
    "St. Petersburg": ("Санкт-Петербург", "Россия"),
    "St.Petersburg": ("Санкт-Петербург", "Россия"),
    "Yekaterinburg": ("Екатеринбург", "Россия"),
}

_BORDER_MAP = {
    "Alashankou": "Алашанькоу",
    "Erlian": "Эрлян",
    "Manzhouli": "Маньчжурия",
}


def _parse_price(val: str) -> float | None:
    """Extract numeric price from string like '$9,900'."""
    if not val:
        return None
    s = val.strip().replace("$", "").replace(",", "").replace(" ", "")
    try:
        return float(s)
    except ValueError:
        return None


def _parse_duration(val: str) -> tuple[int | None, int | None]:
    """Extract duration range from string like '13-15'."""
    if not val:
        return None, None
    m = re.match(r"(\d+)\s*-\s*(\d+)", val.strip())
    if m:
        return int(m.group(1)), int(m.group(2))
    try:
        v = int(val.strip())
        return v, v
    except ValueError:
        return None, None


def _make_segment(
    *,
    start_point: str,
    end_point: str,
    start_country: str,
    end_country: str,
    cost: float,
    currency: str,
    container_type: str,
    border_point: str | None = None,
    duration_min: int | None = None,
    duration_max: int | None = None,
    conditions: str | None = None,
) -> dict:
    return TariffSegment(
        transport_type="truck",
        start_point=start_point,
        end_point=end_point,
        container_type=container_type,
        cost=cost,
        currency=currency,
        company=_COMPANY,
        border_point=border_point,
        duration_min_days=duration_min,
        duration_max_days=duration_max,
        conditions=conditions,
        start_location_type="city",
        end_location_type="city",
        parent_start_location=start_point,
        parent_start_location_type="city",
        parent_end_location=end_point,
        parent_end_location_type="city",
    ).to_dict()


def parse_AmosLogistics(file_path: str | Path | None = None) -> list[dict]:
    """Парсит DOCX с FTL/TIR тарифами Amos Logistics."""
    if file_path is None:
        data_dir = Path(__file__).resolve().parent.parent / "data"
        matches = list(data_dir.glob("Dear Baksanovaig*.docx"))
        if not matches:
            raise FileNotFoundError(
                "Не найден DOCX Amos Logistics в директории data/. "
                "Ожидался файл по шаблону: Dear Baksanovaig*.docx"
            )
        file_path = matches[0]
    else:
        file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    doc = docx.Document(str(file_path))
    results: list[dict] = []

    # ── Table 0: FTL 13.6M tilt tautliner ──
    # Columns: Loading City | Destination | T/F Alashankou | Border | T/F Erlian | Border | T/F Manzhouli | Border
    ftl_table = doc.tables[0]
    for row in ftl_table.rows[2:]:  # skip header rows
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 8:
            continue

        loading_raw = cells[0]
        dest_raw = cells[1]

        if not loading_raw or not dest_raw:
            continue

        start_city, start_country = _LOADING_CITY_MAP.get(
            loading_raw, (loading_raw, "Китай")
        )
        end_city, end_country = _DESTINATION_MAP.get(
            dest_raw, (dest_raw, "Россия")
        )

        start_point = start_city
        end_point = end_city

        # Three border routes: Alashankou (col 2), Erlian (col 4), Manzhouli (col 6)
        border_routes = [
            (cells[2], cells[3]),  # Alashankou
            (cells[4], cells[5]),  # Erlian
            (cells[6], cells[7]),  # Manzhouli
        ]

        for price_raw, border_raw in border_routes:
            price = _parse_price(price_raw)
            if price is None:
                continue

            border = _BORDER_MAP.get(border_raw, border_raw) if border_raw else None

            results.append(_make_segment(
                start_point=start_point,
                end_point=end_point,
                start_country=start_country,
                end_country=end_country,
                cost=price,
                currency="USD",
                container_type="FTL_13.6M",
                border_point=border,
            ))

    # ── Table 1: TIR-Box Truck ──
    # Columns: POL | POD | EXW(RMB) | Border | TT
    tir_table = doc.tables[1]
    for row in tir_table.rows[2:]:  # skip header rows
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue

        pol_raw = cells[0]
        pod_raw = cells[1]

        if not pol_raw or not pod_raw:
            continue

        start_city, start_country = _LOADING_CITY_MAP.get(
            pol_raw, (pol_raw, "Китай")
        )
        end_city, end_country = _DESTINATION_MAP.get(
            pod_raw, (pod_raw, "Россия")
        )

        start_point = start_city
        end_point = end_city

        price = _parse_price(cells[2])
        border = _BORDER_MAP.get(cells[3], cells[3]) if cells[3] else None
        dur_min, dur_max = _parse_duration(cells[4])

        if price is None:
            continue

        results.append(_make_segment(
            start_point=start_point,
            end_point=end_point,
            start_country=start_country,
            end_country=end_country,
            cost=price,
            currency="USD",
            container_type="TIR_BOX_22T",
            border_point=border,
            duration_min=dur_min,
            duration_max=dur_max,
            conditions="TIR — без перегрузки на границе, один водитель",
        ))

    return results


def parse(*args, **kwargs) -> list:
    """Обёртка для единообразия с другими парсерами."""
    return _to_segments(parse_AmosLogistics(*args, **kwargs))
