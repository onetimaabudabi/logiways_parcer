"""Парсер «Хасан» (DOCX) — сквозной морской сервис Китай → Владивосток + ЖД Владивосток → РФ.

Структура исходного файла `data/Хасан (1).docx`:
    Table 0 — СКВОЗНОЙ СЕРВИС КИТАЙ – РОССИЯ (морской фрахт, USD, COC FILO).
    Table 1 — ПЕРЕВОЗКА С ТЕРМИНАЛА ПОРТА ДО ПУНКТА ПРИБЫТИЯ (ЖД, RUB).
    Tables 4–7 — отдельные ЖД ставки Владивосток → Москва (RUB).
    Параграфы — период действия ставки и общий транзитный срок.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document

from .models import TariffSegment
from .utils import to_segments as _to_segments
from shared import (
    container_size_dict,
    currency_dict,
    get_city_port,
    get_country,
    port_dict,
)

_COMPANY = "Хасан"

# Порт прибытия в файле указан как «ВМТП/ВМКТ, НАХОДКА».
# По требованию приводим все морские плечи к единому пункту назначения.
_SEA_POD = "Владивосток"
_SEA_POD_CITY = "Владивосток"
_SEA_POD_TERMINALS = "ВМТП / ВМКТ / Находка"

# Города drop-off из текста КП (общие для всех морских ставок).
_DROPOFF_CITIES = ["Москва", "Санкт-Петербург", "Новосибирск", "Екатеринбург"]

# Порты, у которых русское имя из port_dict не совпадает с городом.
_PORT_PARENT_CITY = {
    "Ксинганг": "Тяньзинь",
    "Наньша": "Гуанчжоу",
    "Яньтянь": "Шэньчжэнь",
}

# Нормализация названий пунктов прибытия ЖД к виду из справочников.
_RAIL_DEST_ALIASES = {
    "МОСКВА": "Москва",
    "САНКТ-ПЕТЕРБУРГ": "Санкт-Петербург",
    "СПБ": "Санкт-Петербург",
    "НОВОСИБИРСК": "Новосибирск",
    "ЕКАТЕРИНБУРГ": "Екатеринбург",
    "ЕКБ": "Екатеринбург",
}

# Весовые категории ЖД-таблицы:
# (container_type, min_weight, max_weight, номер колонки со ставкой)
_RAIL_WEIGHT_CATS = [
    ("20DC", None, 24, 2),
    ("20DC", 24, 28, 3),
    ("40HC", None, 28, 4),
]


def _cell(cell) -> str:
    """Текст ячейки без неразрывных пробелов (переносы строк сохраняем)."""
    return cell.text.replace("\xa0", " ").strip()


def _flat(cell) -> str:
    """Текст ячейки в одну строку."""
    return re.sub(r"\s+", " ", _cell(cell).replace("\n", " ")).strip()


def _table_text(table) -> str:
    return " ".join(_flat(c) for row in table.rows for c in row.cells)


def _parse_price(val) -> Optional[float]:
    """'190 000 ₽' → 190000.0, '2500$' → 2500.0."""
    if val is None:
        return None
    s = re.sub(r"[^\d]", "", str(val).replace("\xa0", ""))
    return float(s) if s else None


def _translate_port(raw: str) -> str:
    """Английское название порта → русское через shared.port_dict.

    'TIANJIN (XINGANG)' → 'Ксинганг'   (прямое совпадение в словаре)
    'GUANGZHOU (NANSHA)' → 'Гуанчжоу'  (по базовому имени без скобок)
    """
    raw = re.sub(r"\*+", "", raw).strip()
    for key in (raw, raw.title(), raw.upper()):
        if key in port_dict:
            return port_dict[key]
    base = re.sub(r"\s*\(.*?\)", "", raw).strip()
    for key in (base, base.title(), base.upper()):
        if key in port_dict:
            return port_dict[key]
    return base.title() or raw.title()


def _split_ports(raw: str) -> list[str]:
    """Разбивает ячейку с несколькими портами.

    'NINGBO / SHANGHAI / QINGDAO' → ['NINGBO', 'SHANGHAI', 'QINGDAO']
    'GUANGZHOU (NANSHA)\\nShenzhen (Yantian, SHEKOU)\\nXIAMEN' → 3 порта
    Разделитель '/' внутри скобок не учитывается.
    """
    parts: list[str] = []
    for line in raw.split("\n"):
        buf, depth = "", 0
        for ch in line:
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth = max(0, depth - 1)
            if ch == "/" and depth == 0:
                parts.append(buf)
                buf = ""
            else:
                buf += ch
        parts.append(buf)
    return [re.sub(r"\s+", " ", p).strip() for p in parts if p.strip()]


def _extract_valid_dates(doc) -> tuple[Optional[str], Optional[str]]:
    """Из текста: 'с 13.01.26 на все рейсы с выходом до 31.01.26' → ('2026-01-13', '2026-01-31')."""
    text = " ".join(p.text for p in doc.paragraphs).replace("\xa0", " ")

    def _norm(d: str, mo: str, y: str) -> str:
        if len(y) == 2:
            y = "20" + y
        return f"{y}-{mo}-{d}"

    valid_from = valid_to = None
    m_from = re.search(r"с\s+(\d{2})\.(\d{2})\.(\d{2,4})", text)
    if m_from:
        valid_from = _norm(*m_from.groups())
    m_to = re.search(r"(?:до|по)\s+(\d{2})\.(\d{2})\.(\d{2,4})", text)
    if m_to:
        valid_to = _norm(*m_to.groups())
    return valid_from, valid_to


def _extract_transit_days(doc) -> Optional[int]:
    """'Общий Транзитный срок доставки ~37 дней.' → 37."""
    text = " ".join(p.text for p in doc.paragraphs).replace("\xa0", " ")
    m = re.search(r"транзитн\w*\s+срок[^.\d]*~?\s*(\d{1,3})\s*дн", text, re.IGNORECASE)
    return int(m.group(1)) if m else None


def _find_table(doc, *keywords):
    """Первая таблица, текст которой содержит все ключевые слова."""
    for table in doc.tables:
        text = _table_text(table).lower()
        if all(kw.lower() in text for kw in keywords):
            return table
    return None


# ────────────────────────────── Морские сегменты ──────────────────────────────


def _parse_sea(doc, valid_from, valid_to, transit_days) -> list[dict]:
    """Table 0 — сквозной сервис Китай → Владивосток (USD)."""
    table = _find_table(doc, "СКВОЗНОЙ СЕРВИС", "Порт отгрузки")
    if table is None:
        return []

    results: list[dict] = []
    seen: set[tuple] = set()
    dropoff = ", ".join(_DROPOFF_CITIES)

    for row in table.rows:
        cells = [_cell(c) for c in row.cells]
        if len(cells) < 4:
            continue

        pol_raw, pod_raw = cells[0], cells[1]
        if not pol_raw or not pod_raw:
            continue
        # Заголовочные строки таблицы
        if re.search(r"порт отгрузки|сквозной сервис|ставка на тип", pol_raw, re.I):
            continue

        price_20 = _parse_price(cells[2])
        price_40 = _parse_price(cells[3])
        if price_20 is None and price_40 is None:
            continue

        # Объединённые ячейки в DOCX дублируют строки — убираем повторы
        key = (pol_raw, pod_raw, price_20, price_40)
        if key in seen:
            continue
        seen.add(key)

        # COC/FILO из заголовка колонки («20DC (СОС FILO)», СОС — кириллица)
        header = f"{cells[2]} {cells[3]}".upper()
        ownership = "SOC" if re.search(r"\bSOC\b", header) else "COC"
        service_term = next(
            (t for t in ("FILO", "FIFO", "LIFO", "LILO") if t in header), "FILO"
        )

        for port_raw in _split_ports(pol_raw):
            port_ru = _translate_port(port_raw)
            start_country = get_country(port_ru) or "Китай"
            parent_city = (
                get_city_port(port_ru) or _PORT_PARENT_CITY.get(port_ru) or port_ru
            )

            for container_type, cost in (("20DC", price_20), ("40HC", price_40)):
                if cost is None:
                    continue
                results.append(
                    TariffSegment(
                        transport_type="sea",
                        start_point=f"{port_ru}, {start_country}",
                        end_point=f"{_SEA_POD}, Россия",
                        container_type=container_type,
                        weight_limit=container_size_dict.get(container_type),
                        max_weight_kg=container_size_dict.get(container_type),
                        cost=cost,
                        currency=currency_dict.get("$", "USD"),
                        company=_COMPANY,
                        container_ownership=ownership,
                        port_service_term=service_term,
                        conditions=(
                            f"Сквозной сервис Китай–Россия, терминалы прибытия: "
                            f"{_SEA_POD_TERMINALS}; drop-off: {dropoff}; "
                            f"free time 50 суток с момента выгрузки"
                        ),
                        duration_min_days=transit_days,
                        duration_max_days=transit_days,
                        valid_from=valid_from,
                        valid_to=valid_to,
                        start_location_type="port",
                        end_location_type="port",
                        parent_start_location=parent_city,
                        parent_start_location_type="city",
                        parent_end_location=_SEA_POD_CITY,
                        parent_end_location_type="city",
                    ).to_dict()
                )

    return results


# ─────────────────────────────── ЖД сегменты ───────────────────────────────


def _make_rail_segment(
    *,
    dest_city: str,
    container_type: str,
    min_weight: Optional[int],
    max_weight: Optional[int],
    cost: float,
    conditions: str,
    valid_from: Optional[str],
    valid_to: Optional[str],
) -> dict:
    weight_limit = (
        container_size_dict.get("20DC_28", "28")
        if container_type == "20DC" and min_weight
        else container_size_dict.get(container_type)
    )
    return TariffSegment(
        transport_type="rail",
        start_point="Владивосток, Россия",
        end_point=f"{dest_city}, {get_country(dest_city) or 'Россия'}",
        container_type=container_type,
        weight_limit=weight_limit,
        min_weight_kg=min_weight,
        max_weight_kg=max_weight,
        cost=cost,
        currency=currency_dict.get("руб", "RUB"),
        company=_COMPANY,
        container_ownership="COC",
        conditions=conditions,
        valid_from=valid_from,
        valid_to=valid_to,
        start_location_type="port",
        end_location_type="rail_station",
        parent_start_location="Владивосток",
        parent_start_location_type="city",
        parent_end_location=dest_city,
        parent_end_location_type="city",
    ).to_dict()


def _parse_rail(doc, valid_from, valid_to) -> list[dict]:
    """Table 1 — ЖД от терминала порта Владивосток до пункта прибытия (RUB)."""
    table = _find_table(doc, "ПЕРЕВОЗКА С ТЕРМИНАЛА ПОРТА", "Пункт отправления")
    if table is None:
        return []

    results: list[dict] = []
    seen: set[tuple] = set()
    conditions = (
        "ЖД от терминала порта Владивосток, НДС 0%; неопасный груз, "
        "не более 1500 кг/место; ставка фиксируется на момент отправки по ЖД"
    )

    for row in table.rows:
        cells = [_flat(c) for c in row.cells]
        if len(cells) < 5:
            continue

        origin_raw, dest_raw = cells[0], cells[1]
        if not origin_raw or not dest_raw or origin_raw == dest_raw:
            continue
        if re.search(r"пункт отправления|перевозка с терминала|ставка", origin_raw, re.I):
            continue

        dest_city = _RAIL_DEST_ALIASES.get(dest_raw.upper(), dest_raw.title())
        if dest_city in seen:
            continue
        seen.add(dest_city)

        for container_type, min_w, max_w, col in _RAIL_WEIGHT_CATS:
            cost = _parse_price(cells[col]) if len(cells) > col else None
            if cost is None:
                continue
            results.append(
                _make_rail_segment(
                    dest_city=dest_city,
                    container_type=container_type,
                    min_weight=min_w,
                    max_weight=max_w,
                    cost=cost,
                    conditions=conditions,
                    valid_from=valid_from,
                    valid_to=valid_to,
                )
            )

    return results


def _parse_rail_moscow(doc, valid_from, valid_to) -> list[dict]:
    """Отдельные ЖД ставки на Москву (Tables 4–7).

    Строки вида: 'Москва 40HC до 28 тн груза брутто' | '278 000 руб., включая НДС-0%'
    """
    # Описание маршрута берём из одной ячейки: в объединённых ячейках DOCX
    # один и тот же текст повторяется по всем колонкам.
    route = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _flat(cell)
                if text.startswith("ЖД перевозка по маршруту"):
                    route = text
                    break
            if route:
                break
        if route:
            break

    conditions = (
        f"Отдельная ЖД ставка, НДС 0%"
        + (f"; {route}" if route else "")
        + "; включено: ПРР в порту, экспедирование, автоуслуги порт – ст. отправления"
    )

    results: list[dict] = []
    seen: set[tuple] = set()

    for table in doc.tables:
        for row in table.rows:
            cells = [_flat(c) for c in row.cells]
            if len(cells) < 2:
                continue
            label, value = cells[0], cells[1]
            if label == value:
                continue

            m = re.match(
                r"^(?P<city>[А-ЯЁ][а-яё\-]+)\s+(?P<ct>20DC|40HC|20GP|40HQ)\s+"
                r"(?P<limit>до|более)\s+(?P<tons>\d{1,3})\s*тн",
                label,
                re.IGNORECASE,
            )
            if not m or "руб" not in value.lower():
                continue

            # «278 000 руб., включая НДС-0%» — отсекаем хвост, иначе 0 из «НДС-0»
            # приклеится к сумме.
            cost = _parse_price(re.split(r"руб", value, maxsplit=1, flags=re.I)[0])
            if cost is None:
                continue

            container_type = m.group("ct").upper()
            tons = int(m.group("tons"))
            if m.group("limit").lower() == "до":
                min_w, max_w = None, tons
            else:
                min_w = tons
                max_w = int(container_size_dict.get("20DC_28", "28"))

            dest_city = _RAIL_DEST_ALIASES.get(
                m.group("city").upper(), m.group("city").title()
            )

            key = (dest_city, container_type, min_w, max_w, cost)
            if key in seen:
                continue
            seen.add(key)

            results.append(
                _make_rail_segment(
                    dest_city=dest_city,
                    container_type=container_type,
                    min_weight=min_w,
                    max_weight=max_w,
                    cost=cost,
                    conditions=conditions,
                    valid_from=valid_from,
                    valid_to=valid_to,
                )
            )

    return results


# ──────────────────────────────── Точка входа ────────────────────────────────


def parse_Khasan_docx(file_path: str | Path | None = None) -> list[dict]:
    """Парсит DOCX «Хасан» и возвращает список словарей-сегментов."""
    if file_path is None:
        data_dir = Path(__file__).resolve().parent.parent / "data"
        matches = sorted(data_dir.glob("Хасан*.docx"))
        if not matches:
            raise FileNotFoundError(
                "Не найден DOCX Хасан в директории data/. "
                "Ожидался файл по шаблону: Хасан*.docx"
            )
        file_path = matches[0]
    else:
        file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    doc = Document(str(file_path))

    valid_from, valid_to = _extract_valid_dates(doc)
    transit_days = _extract_transit_days(doc)

    results: list[dict] = []
    results += _parse_sea(doc, valid_from, valid_to, transit_days)
    results += _parse_rail(doc, valid_from, valid_to)
    results += _parse_rail_moscow(doc, valid_from, valid_to)
    return results


def parse(*args, **kwargs) -> list[TariffSegment]:
    """Обёртка для единообразия с другими парсерами."""
    return _to_segments(parse_Khasan_docx(*args, **kwargs))
